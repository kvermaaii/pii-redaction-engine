import sys
from docx import Document
from faker import Faker
from presidio_analyzer import AnalyzerEngine, PatternRecognizer, Pattern
from presidio_analyzer.nlp_engine import NlpEngineProvider
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig

ALLOW_LIST = [
    "E-MAIL", "EMAIL", "E-Mail", "Email", "Email:", "E-Mail:", "E-MAIL AND TELEPHONE",
    "TELEPHONE", "Telephone", "Telephone:", "PHONE", "Phone", "Phone:",
    "FAX", "Fax", "Fax:", "WEBSITE", "Website", "Website:",
    "CONTACT PERSON", "Contact Person", "REGISTERED OFFICE", "Registered Office",
    "CORPORATE OFFICE", "Corporate Office", "OUR PROMOTERS", "PROMOTERS", "Promoters",
    "DETAILS OF THE OFFER TO PUBLIC", "TYPE", "SIZE OF THE FRESH ISSUE",
    "SIZE OF THE OFFER FOR SALE", "TOTAL OFFER SIZE", "ELIGIBILITY",
    "RESERVATION AMONG QIBs, NIIs AND RIIs", "QIBs", "NIIs", "RIIs",
    "Ticket", "Order", "Application", "Section", "Clause", "Table",
    "Red Herring Prospectus", "Prospectus", "Draft Red Herring Prospectus"
]

def setup_engines(score_threshold=0.4):
    """Initializes the Presidio Analyzer with custom recognizers, Anonymizer, and Faker."""
    configuration = {
        "nlp_engine_name": "spacy",
        "models": [{"lang_code": "en", "model_name": "en_core_web_sm"}],
    }
    provider = NlpEngineProvider(nlp_configuration=configuration)
    nlp_engine = provider.create_engine()
    analyzer = AnalyzerEngine(nlp_engine=nlp_engine, supported_languages=["en"])
    
    pan_pattern = Pattern(name="pan_pattern", regex=r"\b[A-Z]{5}[0-9]{4}[A-Z]{1}\b", score=0.95)
    pan_recognizer = PatternRecognizer(supported_entity="US_SSN", patterns=[pan_pattern])
    
    aadhaar_pattern = Pattern(name="aadhaar_pattern", regex=r"\b[2-9]{1}[0-9]{3}\s?[0-9]{4}\s?[0-9]{4}\b", score=0.90)
    aadhaar_recognizer = PatternRecognizer(supported_entity="US_SSN", patterns=[aadhaar_pattern])
    
    in_phone_pattern = Pattern(name="in_phone_pattern", regex=r"\b(?:\+91[\-\s]?)?[6-9]\d{9}\b|\b0\d{2,4}[\-\s]?\d{6,8}\b", score=0.85)
    in_phone_recognizer = PatternRecognizer(supported_entity="PHONE_NUMBER", patterns=[in_phone_pattern])
    
    dob_pattern = Pattern(name="dob_pattern", regex=r"\b(?:0[1-9]|[12][0-9]|3[01])[\/\.-](?:0[1-9]|1[012])[\/\.-](?:19|20)\d{2}\b|\b(?:19|20)\d{2}[\/\.-](?:0[1-9]|1[012])[\/\.-](?:0[1-9]|[12][0-9]|3[01])\b", score=0.85)
    dob_recognizer = PatternRecognizer(supported_entity="DATE_TIME", patterns=[dob_pattern])

    analyzer.registry.add_recognizer(pan_recognizer)
    analyzer.registry.add_recognizer(aadhaar_recognizer)
    analyzer.registry.add_recognizer(in_phone_recognizer)
    analyzer.registry.add_recognizer(dob_recognizer)

    anonymizer = AnonymizerEngine()
    fake = Faker()

    operators = {
        "PERSON": OperatorConfig("custom", {"lambda": lambda x: fake.name()}),
        "EMAIL_ADDRESS": OperatorConfig("custom", {"lambda": lambda x: fake.email()}),
        "PHONE_NUMBER": OperatorConfig("custom", {"lambda": lambda x: fake.phone_number()}),
        "ORGANIZATION": OperatorConfig("custom", {"lambda": lambda x: fake.company()}),
        "LOCATION": OperatorConfig("custom", {"lambda": lambda x: fake.address().replace('\n', ', ')}),
        "US_SSN": OperatorConfig("custom", {"lambda": lambda x: fake.ssn()}),
        "CREDIT_CARD": OperatorConfig("custom", {"lambda": lambda x: fake.credit_card_number()}),
        "DATE_TIME": OperatorConfig("custom", {"lambda": lambda x: fake.date_of_birth().strftime('%Y-%m-%d')}), 
        "IP_ADDRESS": OperatorConfig("custom", {"lambda": lambda x: fake.ipv4()})
    }
    
    return analyzer, anonymizer, operators, score_threshold

def redact_text(text: str, analyzer, anonymizer, operators, score_threshold=0.4) -> str:
    if not text.strip():
        return text
    
    entities = list(operators.keys())
    results = analyzer.analyze(
        text=text,
        entities=entities,
        language='en',
        score_threshold=score_threshold,
        allow_list=ALLOW_LIST
    )
    
    if not results:
        return text
        
    anonymized_result = anonymizer.anonymize(
        text=text,
        analyzer_results=results,
        operators=operators
    )
    return anonymized_result.text

def process_docx(input_path: str, output_path: str, score_threshold=0.4):
    analyzer, anonymizer, operators, threshold = setup_engines(score_threshold=score_threshold)
    doc = Document(input_path)
    
    redacted_count = 0
    for para in doc.paragraphs:
        if para.text.strip():
            old_text = para.text
            new_text = redact_text(old_text, analyzer, anonymizer, operators, score_threshold=threshold)
            if old_text != new_text:
                redacted_count += 1
            para.text = new_text
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        old_text = para.text
                        new_text = redact_text(old_text, analyzer, anonymizer, operators, score_threshold=threshold)
                        if old_text != new_text:
                            redacted_count += 1
                        para.text = new_text
                        
    try:
        doc.save(output_path)
        print(f"Redaction complete! Total redacted elements: {redacted_count}")
    except PermissionError:
        fallback_path = output_path.replace(".docx", "_Final.docx")
        doc.save(fallback_path)
        print(f"Redaction complete! Total redacted elements: {redacted_count}")

if __name__ == "__main__":
    input_file = sys.argv[1] if len(sys.argv) > 1 else "Red Herring Prospectus.docx"
    output_file = sys.argv[2] if len(sys.argv) > 2 else "Redacted_Output_Final_Final.docx"
    process_docx(input_file, output_file)
